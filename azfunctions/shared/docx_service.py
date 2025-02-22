import io
from docx import Document
from typing import Dict, List, Optional

from shared.models import DocumentPage, DocumentStyle, Line, TableCell


class DocxService:
    @staticmethod
    def get_text_from_docx(document_content) -> dict:
        doc = Document(io.BytesIO(document_content))
        full_text = []
        paragraphs = []
        tables = []
        styles = {}
        headers = []
        footers = []
        pages = []  # Note: python-docx doesn't provide direct page information
        
        # Extract styles
        for style in doc.styles:
            if hasattr(style, 'name') and style.name:
                font = style.font if hasattr(style, 'font') else None
                styles[style.name] = DocumentStyle(
                    name=style.name,
                    font_name=font.name if font and hasattr(font, 'name') else None,
                    font_size=float(font.size.pt) if font and hasattr(font, 'size') and font.size else None,
                    is_bold=font.bold if font and hasattr(font, 'bold') else None,
                    is_italic=font.italic if font and hasattr(font, 'italic') else None,
                    is_underline=font.underline if font and hasattr(font, 'underline') else None
                )
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                paragraphs.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(TableCell(text=cell.text))
                    full_text.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
        
        # Extract text from text boxes and shapes
        for shape in doc.inline_shapes:
            if shape.text_frame:
                for paragraph in shape.text_frame.paragraphs:
                    if paragraph.text.strip():
                        full_text.append(paragraph.text)
                        paragraphs.append(paragraph.text)
        
        if hasattr(doc, 'shapes'):
            for shape in doc.shapes:
                if shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
                            paragraphs.append(paragraph.text)
        
        # Extract headers and footers from sections
        for section in doc.sections:
            if section.header:
                for paragraph in section.header.paragraphs:
                    header_text = paragraph.text.strip()
                    if header_text:
                        headers.append(header_text)
                        full_text.append(header_text)
            
            if section.footer:
                for paragraph in section.footer.paragraphs:
                    footer_text = paragraph.text.strip()
                    if footer_text:
                        footers.append(footer_text)
                        full_text.append(footer_text)
        
        # Create a single page since python-docx doesn't provide page information
        pages.append(DocumentPage(
            page_number=1,
            content='\n'.join(full_text),
            lines=[Line(content=text) for text in full_text],
            tables=tables
        ))
        
        # Create structured document info
        structured_info = {
            'text': '\n'.join(full_text),
            'pages': pages,
            'paragraphs': paragraphs,
            'tables': tables,
            'styles': styles,
            'headers': headers,
            'footers': footers
        }
        
        return structured_info

    # Replace 'your_document.docx' with the path to your Word document
    # document_path = 'your_document.docx'
    # document_text = get_text_from_docx(document_path)
    # print(document_text)
